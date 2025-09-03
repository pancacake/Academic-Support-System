import os
import json
import threading
from django.http import JsonResponse, StreamingHttpResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from django.conf import settings
from .note_generator import NoteGenerator

# 导入集中管理的提示词
try:
    from prompts import CHAT_ASSISTANT_PROMPT
except ImportError:
    CHAT_ASSISTANT_PROMPT = "你是一个学术助手，请基于提供的笔记内容回答用户问题。"

# 全局变量存储笔记生成状态
note_generation_status = {}

def get_user_id(request):
    """获取用户ID，未登录返回0"""
    if request and hasattr(request, 'user') and request.user.is_authenticated:
        return request.user.id
    return 0

def get_user_upload_path(user_id):
    """获取用户上传目录路径"""
    return os.path.join('media', str(user_id), 'uploads')

def get_user_output_path(user_id):
    """获取用户输出目录路径"""
    return os.path.join('media', str(user_id), 'output')

@csrf_exempt
@api_view(['POST'])
def start_note_generation(request):
    """开始笔记生成"""
    try:
        user_id = get_user_id(request)
        print(f"[DEBUG] start_note_generation - 用户ID: {user_id}")
        print(f"[DEBUG] start_note_generation - 请求用户: {getattr(request, 'user', 'No user')}")
        print(f"[DEBUG] start_note_generation - 用户认证: {getattr(request.user, 'is_authenticated', 'No auth') if hasattr(request, 'user') else 'No user attr'}")

        # 获取上传目录中的所有已解析文件
        upload_dir = get_user_upload_path(user_id)
        print(f"[DEBUG] start_note_generation - 上传目录: {upload_dir}")
        if not os.path.exists(upload_dir):
            return JsonResponse({'success': False, 'error': '没有找到上传的文件'}, status=400)

        # 查找所有解析后的JSON文件
        json_files = []
        for item in os.listdir(upload_dir):
            item_path = os.path.join(upload_dir, item)
            if os.path.isdir(item_path):
                # 查找目录中的JSON文件
                for file in os.listdir(item_path):
                    if file.endswith('.json'):
                        json_path = os.path.join(item_path, file)
                        json_files.append({
                            'name': file,
                            'path': json_path,
                            'folder': item
                        })

        if not json_files:
            return JsonResponse({'success': False, 'error': '没有找到已解析的文件'}, status=400)

        # 设置生成状态
        note_generation_status['current'] = {
            'status': 'preparing',
            'message': '正在准备生成笔记...',
            'files': json_files
        }

        # 异步生成笔记
        def generate_notes():
            try:
                generator = NoteGenerator()

                # 合并所有JSON文件的内容
                all_content = []
                for json_file in json_files:
                    try:
                        with open(json_file['path'], 'r', encoding='utf-8') as f:
                            content = json.load(f)
                            if isinstance(content, list):
                                all_content.extend(content)
                    except Exception as e:
                        print(f"读取文件 {json_file['path']} 失败: {e}")

                if not all_content:
                    note_generation_status['current'] = {
                        'status': 'error',
                        'message': '没有找到有效的文件内容'
                    }
                    return

                # 创建临时合并文件
                temp_json_path = os.path.join(upload_dir, 'merged_content.json')
                with open(temp_json_path, 'w', encoding='utf-8') as f:
                    json.dump(all_content, f, ensure_ascii=False, indent=2)

                # 开始流式生成
                note_generation_status['current'] = {
                    'status': 'generating',
                    'message': '正在生成笔记...',
                    'content': ''
                }

                print("[DEBUG] 开始调用笔记生成器")  # 调试信息

                # 使用用户特定的输出目录
                user_output_dir = get_user_output_path(user_id)

                for chunk in generator.generate_notes_streaming(temp_json_path, user_output_dir):
                    print(f"[DEBUG] 收到chunk: {chunk['type']}")  # 调试信息

                    if chunk['type'] == 'start':
                        note_generation_status['current'] = {
                            'status': 'generating',
                            'message': chunk['content'],
                            'content': ''
                        }
                    elif chunk['type'] == 'content':
                        # 累积内容
                        current_status = note_generation_status.get('current', {})
                        current_content = current_status.get('content', '')
                        current_content += chunk['content']

                        note_generation_status['current'] = {
                            'status': 'generating',
                            'message': '正在生成笔记...',
                            'content': current_content,
                            'chunk': chunk['content']  # 新增的内容块
                        }
                        print(f"[DEBUG] 内容长度: {len(current_content)}")  # 调试信息

                    elif chunk['type'] == 'complete':
                        final_content = note_generation_status.get('current', {}).get('content', '')
                        note_generation_status['current'] = {
                            'status': 'completed',
                            'message': '笔记生成完成！您可以在右侧查看生成的笔记。有什么问题吗？',
                            'content': final_content,
                            'file_path': chunk['file_path'],
                            'output_dir': chunk['output_dir']
                        }
                        print(f"[DEBUG] 笔记生成完成，最终内容长度: {len(final_content)}")
                        break

                    elif chunk['type'] == 'error':
                        note_generation_status['current'] = {
                            'status': 'error',
                            'message': f'笔记生成失败：{chunk["content"]}'
                        }
                        print(f"[DEBUG] 笔记生成错误: {chunk['content']}")
                        break

                # 清理临时文件
                try:
                    os.remove(temp_json_path)
                except:
                    pass

            except Exception as e:
                note_generation_status['current'] = {
                    'status': 'error',
                    'message': f'笔记生成失败：{str(e)}'
                }

        thread = threading.Thread(target=generate_notes)
        thread.daemon = True
        thread.start()

        return JsonResponse({
            'success': True,
            'message': f'找到 {len(json_files)} 个已解析的文件，正在开始生成笔记...',
            'files': [f['name'] for f in json_files]
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': f'启动笔记生成失败：{str(e)}'}, status=500)

@api_view(['GET'])
def get_note_generation_status(request):
    """获取笔记生成状态"""
    if 'current' not in note_generation_status:
        return Response({'status': 'none', 'message': '暂无笔记生成任务。'})
    return Response(note_generation_status['current'])

@csrf_exempt
def stream_notes_content(request):
    """流式传输笔记内容 - 直接生成模式"""
    print(f"[DEBUG] 收到流式请求: {request.method} {request.path}")

    # 在生成器外部获取用户信息
    user_id = get_user_id(request)
    upload_dir = get_user_upload_path(user_id)
    print(f"[DEBUG] stream_notes - 用户ID: {user_id}")
    print(f"[DEBUG] stream_notes - 上传目录: {upload_dir}")
    print(f"[DEBUG] stream_notes - 目录是否存在: {os.path.exists(upload_dir)}")

    def generate():
        try:
            # 发送连接成功消息
            yield "data: " + json.dumps({'type': 'connected', 'message': '连接成功，开始生成笔记...'}, ensure_ascii=False) + "\n\n"

            if not os.path.exists(upload_dir):
                yield "data: " + json.dumps({'type': 'error', 'message': '没有找到上传的文件'}, ensure_ascii=False) + "\n\n"
                return

            # 查找所有解析后的JSON文件
            json_files = []
            for item in os.listdir(upload_dir):
                item_path = os.path.join(upload_dir, item)
                if os.path.isdir(item_path):
                    for file in os.listdir(item_path):
                        if file.endswith('.json'):
                            json_path = os.path.join(item_path, file)
                            json_files.append(json_path)
                            print(f"[DEBUG] stream_notes - 找到JSON文件: {json_path}")

            print(f"[DEBUG] stream_notes - 总共找到 {len(json_files)} 个JSON文件")
            if not json_files:
                yield "data: " + json.dumps({'type': 'error', 'message': '没有找到已解析的文件'}, ensure_ascii=False) + "\n\n"
                return

            yield "data: " + json.dumps({'type': 'preparing', 'message': f'找到 {len(json_files)} 个文件，正在准备生成...'}, ensure_ascii=False) + "\n\n"

            # 合并所有JSON文件的内容
            all_content = []
            for json_file in json_files:
                try:
                    with open(json_file, 'r', encoding='utf-8') as f:
                        content = json.load(f)
                        if isinstance(content, list):
                            all_content.extend(content)
                except Exception as e:
                    print(f"读取文件 {json_file} 失败: {e}")

            if not all_content:
                yield "data: " + json.dumps({'type': 'error', 'message': '没有找到有效的文件内容'}, ensure_ascii=False) + "\n\n"
                return

            # 创建临时合并文件
            temp_json_path = os.path.join(upload_dir, 'merged_content.json')
            with open(temp_json_path, 'w', encoding='utf-8') as f:
                json.dump(all_content, f, ensure_ascii=False, indent=2)

            # 直接开始流式生成
            from .note_generator import NoteGenerator
            generator = NoteGenerator()

            yield "data: " + json.dumps({'type': 'start', 'message': '开始生成笔记，请耐心等待（可能需要1-2分钟）...'}, ensure_ascii=False) + "\n\n"

            # 简化实现，先不使用心跳线程
            print("[DEBUG] 开始笔记生成流程")

            try:
                # 使用用户特定的输出目录
                user_output_dir = get_user_output_path(user_id)

                for chunk in generator.generate_notes_streaming(temp_json_path, user_output_dir):
                    if chunk['type'] == 'start':
                        yield "data: " + json.dumps({'type': 'start', 'message': chunk['content']}, ensure_ascii=False) + "\n\n"
                    elif chunk['type'] == 'content':
                        yield "data: " + json.dumps({'type': 'content', 'content': chunk['content']}, ensure_ascii=False) + "\n\n"
                    elif chunk['type'] == 'complete':
                        yield "data: " + json.dumps({
                            'type': 'complete',
                            'message': chunk['content'],
                            'file_path': chunk.get('file_path', ''),
                            'output_dir': chunk.get('output_dir', ''),
                            'toc_content': chunk.get('toc_content', ''),
                            'toc_file_path': chunk.get('toc_file_path', '')
                        }, ensure_ascii=False) + "\n\n"
                        break
                    elif chunk['type'] == 'error':
                        yield "data: " + json.dumps({'type': 'error', 'message': chunk['content']}, ensure_ascii=False) + "\n\n"
                        break
            finally:
                print("[DEBUG] 笔记生成流程结束")

            # 清理临时文件
            try:
                os.remove(temp_json_path)
            except:
                pass

        except Exception as e:
            print(f"[ERROR] 流式生成错误: {e}")
            yield "data: " + json.dumps({'type': 'error', 'message': f'生成错误: {str(e)}'}, ensure_ascii=False) + "\n\n"

    response = StreamingHttpResponse(generate(), content_type='text/event-stream; charset=utf-8')
    response['Cache-Control'] = 'no-cache'
    # 移除Connection头部，WSGI不允许设置hop-by-hop头部
    response['Access-Control-Allow-Origin'] = '*'
    response['Access-Control-Allow-Headers'] = 'Cache-Control'
    # 设置更长的超时时间
    response['X-Accel-Buffering'] = 'no'  # 禁用nginx缓冲
    return response

@csrf_exempt
def simple_stream_test(request):
    """简单的流式测试"""
    print(f"[DEBUG] 简单流式测试: {request.method} {request.path}")

    def simple_generate():
        import time
        import json

        try:
            yield "data: " + json.dumps({'type': 'start', 'message': '测试开始'}, ensure_ascii=False) + "\n\n"

            for i in range(5):
                time.sleep(1)
                yield "data: " + json.dumps({'type': 'content', 'content': f'测试内容 {i+1}'}, ensure_ascii=False) + "\n\n"

            yield "data: " + json.dumps({'type': 'complete', 'message': '测试完成'}, ensure_ascii=False) + "\n\n"

        except Exception as e:
            print(f"[ERROR] 简单流式测试错误: {e}")
            yield "data: " + json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False) + "\n\n"

    response = StreamingHttpResponse(simple_generate(), content_type='text/event-stream; charset=utf-8')
    response['Cache-Control'] = 'no-cache'
    # 移除Connection头部，WSGI不允许设置hop-by-hop头部
    response['Access-Control-Allow-Origin'] = '*'
    return response

@csrf_exempt
@api_view(['POST'])
def ai_chat_with_notes(request):
    """AI对话功能，支持笔记内容修改"""
    try:
        data = request.data
        user_message = data.get('message', '')

        if not user_message:
            return Response({'success': False, 'error': '消息不能为空'}, status=400)

        # 获取最新的笔记文件
        user_id = get_user_id(request)
        latest_notes = get_latest_notes_file(user_id)
        if not latest_notes:
            return Response({'success': False, 'error': '没有找到笔记文件'}, status=404)

        notes_content = latest_notes['content']
        notes_file_path = latest_notes['file_path']

        # 检查用户是否想要修改特定部分
        section_match = extract_section_from_message(user_message, notes_content)

        if section_match:
            # 用户想要修改特定部分
            section_title = section_match['title']
            section_content = section_match['content']

            # 调用AI生成改进内容
            improved_content = generate_improved_section(user_message, section_title, section_content)

            # 替换笔记中的内容
            updated_notes = replace_section_in_notes(notes_content, section_title, improved_content)

            # 保存更新后的笔记
            with open(notes_file_path, 'w', encoding='utf-8') as f:
                f.write(updated_notes)

            # 重新生成目录
            from .note_generator import NoteGenerator
            generator = NoteGenerator()
            toc_content = generator._generate_table_of_contents(notes_file_path)
            toc_file_path = os.path.join(os.path.dirname(notes_file_path), 'contents.md')
            with open(toc_file_path, 'w', encoding='utf-8') as f:
                f.write(toc_content)

            return Response({
                'success': True,
                'message': f'已成功更新"{section_title}"部分的内容',
                'section_title': section_title,
                'updated_content': improved_content,
                'updated_toc': toc_content
            })
        else:
            # 普通AI对话
            ai_response = generate_ai_response(user_message, notes_content)

            return Response({
                'success': True,
                'message': ai_response,
                'is_general_chat': True
            })

    except Exception as e:
        return Response({'success': False, 'error': f'AI对话失败：{str(e)}'}, status=500)

def get_latest_notes_file(user_id=0):
    """获取最新的笔记文件"""
    try:
        output_dir = get_user_output_path(user_id)
        if not os.path.exists(output_dir):
            return None

        # 获取所有输出目录，按时间排序
        folders = []
        for item in os.listdir(output_dir):
            item_path = os.path.join(output_dir, item)
            if os.path.isdir(item_path):
                notes_file = os.path.join(item_path, 'notes.md')
                if os.path.exists(notes_file):
                    folders.append({
                        'name': item,
                        'path': item_path,
                        'notes_file': notes_file,
                        'created': os.path.getctime(item_path)
                    })

        if not folders:
            return None

        # 按创建时间排序，获取最新的
        latest_folder = sorted(folders, key=lambda x: x['created'], reverse=True)[0]

        # 读取笔记内容
        with open(latest_folder['notes_file'], 'r', encoding='utf-8') as f:
            content = f.read()

        return {
            'folder': latest_folder['name'],
            'content': content,
            'file_path': latest_folder['notes_file']
        }

    except Exception as e:
        print(f"获取最新笔记文件失败: {e}")
        return None

def extract_section_from_message(user_message, notes_content):
    """从用户消息中提取要修改的章节"""
    import re

    # 提取笔记中的所有标题
    lines = notes_content.split('\n')
    headers = []

    for i, line in enumerate(lines):
        line = line.strip()
        if line.startswith('#'):
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break

            title = line[level:].strip()
            if title:
                headers.append({
                    'title': title,
                    'level': level,
                    'line_index': i
                })

    # 尝试在用户消息中找到匹配的标题
    user_message_lower = user_message.lower()

    # 常见的修改意图关键词
    modify_keywords = ['修改', '改进', '优化', '重写', '更新', '完善', '补充', '详细', '简化']

    # 检查是否包含修改意图
    has_modify_intent = any(keyword in user_message_lower for keyword in modify_keywords)

    if not has_modify_intent:
        return None

    # 寻找最匹配的标题
    best_match = None
    best_score = 0

    for header in headers:
        title_lower = header['title'].lower()

        # 计算匹配度
        score = 0
        title_words = title_lower.split()

        for word in title_words:
            if word in user_message_lower:
                score += len(word)

        # 如果用户消息中包含完整的标题，给予更高分数
        if title_lower in user_message_lower:
            score += len(title_lower) * 2

        if score > best_score:
            best_score = score
            best_match = header

    if best_match and best_score > 2:  # 至少要有一定的匹配度
        # 提取该章节的内容
        section_content = extract_section_content(notes_content, best_match['title'])

        return {
            'title': best_match['title'],
            'content': section_content,
            'level': best_match['level']
        }

    return None

def extract_section_content(content, section_title):
    """从笔记中提取特定章节的内容"""
    lines = content.split('\n')
    section_lines = []
    in_section = False
    current_level = 0

    for line in lines:
        line_stripped = line.strip()

        # 检查是否是标题行
        if line_stripped.startswith('#'):
            level = len(line_stripped) - len(line_stripped.lstrip('#'))
            title = line_stripped[level:].strip()

            if title == section_title:
                in_section = True
                current_level = level
                section_lines.append(line)
                continue
            elif in_section and level <= current_level:
                # 遇到同级或更高级标题，结束当前章节
                break

        if in_section:
            section_lines.append(line)

    return '\n'.join(section_lines)

def generate_improved_section(user_message, section_title, section_content):
    """调用AI生成改进的章节内容"""
    try:
        # 导入API客户端
        from .note_generator import get_api_client

        client = get_api_client()
        if not client:
            return f"# {section_title}\n\nAPI客户端不可用，无法生成改进内容。"

        prompt = f"""
用户对笔记中的"{section_title}"部分提出了以下要求：
{user_message}

当前该部分的内容是：
{section_content}

请根据用户的要求，重新生成这个部分的内容。要求：
1. 保持Markdown格式
2. 保持标题级别不变
3. 内容要更加详细、准确、易懂
4. 如果用户有具体要求，请严格按照要求修改
5. 保持学术性和专业性

请直接输出改进后的内容，不需要额外说明：
"""

        response = client.chat_completion([
            {"role": "user", "content": prompt}
        ])

        return response.strip()

    except Exception as e:
        return f"# {section_title}\n\n生成改进内容时出错: {str(e)}"

def replace_section_in_notes(notes_content, section_title, new_content):
    """替换笔记中的特定章节内容"""
    lines = notes_content.split('\n')
    new_lines = []
    in_section = False
    current_level = 0

    for line in lines:
        line_stripped = line.strip()

        # 检查是否是标题行
        if line_stripped.startswith('#'):
            level = len(line_stripped) - len(line_stripped.lstrip('#'))
            title = line_stripped[level:].strip()

            if title == section_title:
                # 找到要替换的章节，添加新内容
                new_lines.extend(new_content.split('\n'))
                in_section = True
                current_level = level
                continue
            elif in_section and level <= current_level:
                # 遇到同级或更高级标题，结束替换，添加当前行
                in_section = False
                new_lines.append(line)
                continue

        if not in_section:
            new_lines.append(line)

    return '\n'.join(new_lines)

def generate_ai_response(user_message, notes_content):
    """生成普通AI对话回复"""
    try:
        from .note_generator import get_api_client

        client = get_api_client()
        if not client:
            return "抱歉，AI服务暂时不可用。"

        # 使用集中管理的提示词
        prompt = CHAT_ASSISTANT_PROMPT.format(
            notes_content=notes_content[:1000] + "..." if len(notes_content) > 1000 else notes_content,
            user_question=user_message
        )

        response = client.chat_completion([
            {"role": "user", "content": prompt}
        ])

        return response.strip()

    except Exception as e:
        return f"抱歉，生成回复时出错: {str(e)}"

def extract_section_content(notes_content, section_title):
    """从笔记内容中提取指定章节的内容"""
    lines = notes_content.split('\n')
    section_lines = []
    in_target_section = False
    current_level = 0

    for line in lines:
        # 检查是否是标题行
        if line.strip().startswith('#'):
            # 计算标题级别
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break

            # 提取标题文本
            title_text = line[level:].strip()

            # 检查是否是目标章节
            if section_title.lower() in title_text.lower() or title_text.lower() in section_title.lower():
                in_target_section = True
                current_level = level
                section_lines.append(line)
            elif in_target_section:
                # 如果遇到同级或更高级的标题，结束当前章节
                if level <= current_level:
                    break
                else:
                    section_lines.append(line)
            else:
                continue
        elif in_target_section:
            section_lines.append(line)

    return '\n'.join(section_lines).strip()

@api_view(['GET'])
def get_note_sections(request):
    """获取笔记章节列表用于@提及"""
    try:
        # 获取最新的笔记文件
        user_id = get_user_id(request)
        latest_notes = get_latest_notes_file(user_id)
        if not latest_notes:
            return Response({'success': False, 'error': '没有找到笔记文件'}, status=404)

        notes_content = latest_notes['content']

        # 解析二级标题
        sections = []
        lines = notes_content.split('\n')

        for i, line in enumerate(lines):
            line = line.strip()
            if not line.startswith('#'):
                continue

            # 计算标题级别
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break

            # 只获取二级标题
            if level == 2:
                title = line[level:].strip()
                if title:
                    # 获取预览内容（接下来的几行非标题内容）
                    preview_lines = []
                    for j in range(i + 1, min(i + 4, len(lines))):
                        preview_line = lines[j].strip()
                        if preview_line and not preview_line.startswith('#'):
                            preview_lines.append(preview_line)
                        elif preview_line.startswith('#'):
                            break

                    preview = ' '.join(preview_lines)[:100] + ('...' if len(' '.join(preview_lines)) > 100 else '')

                    sections.append({
                        'title': title,
                        'preview': preview or '暂无内容预览',
                        'level': level
                    })

        return Response({
            'success': True,
            'sections': sections
        })

    except Exception as e:
        return Response({'success': False, 'error': f'获取章节失败：{str(e)}'}, status=500)

def export_notes(request):
    """导出笔记文件"""
    try:
        format_type = request.GET.get('format', 'md')  # md 或 docx
        print(f"[DEBUG] 导出请求 - 格式: {format_type}")

        # 获取最新的笔记文件
        user_id = get_user_id(request)
        print(f"[DEBUG] 用户ID: {user_id}")

        latest_notes = get_latest_notes_file(user_id)
        print(f"[DEBUG] 最新笔记: {latest_notes}")

        if not latest_notes:
            print("[DEBUG] 没有找到笔记文件")
            return JsonResponse({'success': False, 'error': '没有找到笔记文件。请先生成笔记后再尝试导出。'}, status=404)

        notes_content = latest_notes['content']
        notes_file_path = latest_notes['file_path']

        if format_type == 'md':
            # 导出Markdown格式
            print(f"[DEBUG] 导出MD格式，内容长度: {len(notes_content)}")
            response = HttpResponse(notes_content, content_type='text/markdown; charset=utf-8')
            response['Content-Disposition'] = 'attachment; filename="notes.md"'
            return response

        elif format_type == 'docx':
            # 导出DOCX格式
            print(f"[DEBUG] 导出DOCX格式")
            docx_content = convert_markdown_to_docx(notes_content)

            response = HttpResponse(docx_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="notes.docx"'
            return response

        else:
            return JsonResponse({'success': False, 'error': '不支持的导出格式'}, status=400)

    except Exception as e:
        print(f"[ERROR] 导出失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': f'导出失败：{str(e)}'}, status=500)

def convert_markdown_to_docx(markdown_content):
    """将Markdown内容转换为DOCX格式"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re

        doc = Document()

        # 设置文档标题
        title = doc.add_heading('学习笔记', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        lines = markdown_content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                # 空行
                doc.add_paragraph()
                continue

            # 处理标题
            if line.startswith('#'):
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break

                title_text = line[level:].strip()
                if title_text:
                    doc.add_heading(title_text, level)
                continue

            # 处理列表
            if line.startswith('- ') or line.startswith('* ') or line.startswith('+ '):
                list_text = line[2:].strip()
                p = doc.add_paragraph(list_text, style='List Bullet')
                continue

            # 处理有序列表
            if re.match(r'^\d+\.\s+', line):
                list_text = re.sub(r'^\d+\.\s+', '', line)
                p = doc.add_paragraph(list_text, style='List Number')
                continue

            # 处理代码块
            if line.startswith('```'):
                continue  # 跳过代码块标记

            # 处理普通段落
            if line:
                # 处理粗体和斜体
                line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # 简化处理粗体
                line = re.sub(r'\*([^*]+)\*', r'\1', line)      # 简化处理斜体
                line = re.sub(r'`([^`]+)`', r'\1', line)        # 简化处理行内代码

                doc.add_paragraph(line)

        # 保存到内存
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    except ImportError:
        # 如果没有安装python-docx，返回错误
        raise Exception("需要安装python-docx库才能导出DOCX格式。请运行: pip install python-docx")
    except Exception as e:
        raise Exception(f"转换DOCX格式失败: {str(e)}")
